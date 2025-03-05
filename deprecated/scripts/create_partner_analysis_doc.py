#!/usr/bin/env python3
"""
Create Partner Analysis Summary Document

This script creates a Word document with a summary of the partner analysis script.

Usage:
    python create_partner_analysis_doc.py

Output:
    Creates a Word document in the data/output directory
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # Create output directory if it doesn't exist
    output_dir = "data/output"
    os.makedirs(output_dir, exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Partner Analysis Script Summary"
    doc.core_properties.author = "Dell-AITC"
    
    # Add title
    title = doc.add_heading("Partner Analysis Script Summary", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph("The partner analysis script (partner_analysis.py) is designed to evaluate and categorize partners against AI technology categories. Here's a comprehensive overview of how it works:")
    
    # Core Functionality section
    doc.add_heading("Core Functionality", level=2)
    doc.add_paragraph("The script analyzes partners to determine which AI technology categories they match with by:")
    
    # Add bullet points for Core Functionality
    bullet_points = [
        "Data Collection: Gathering information from multiple sources:",
        "- Bootstrap information (basic partner details)",
        "- Website content",
        "- Search results (via Brave Search API)",
        "- GitHub repositories",
        "AI-Powered Analysis: Using LLM models to evaluate the collected data and match partners to AI technology categories.",
        "Output Generation: Producing structured analysis results in CSV and JSON formats."
    ]
    
    for point in bullet_points:
        if point.startswith("-"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(point[2:])
        else:
            doc.add_paragraph(point, style="List Bullet")
    
    # Workflow Process section
    doc.add_heading("Workflow Process", level=2)
    
    doc.add_paragraph("1. Initialization:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("- Parses command-line arguments\n- Configures logging\n- Loads AI technology categories from a CSV file\n- Loads partner data from a CSV file")
    
    doc.add_paragraph("2. For Each Partner:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("- Bootstrap Information Collection: Gets basic information about the partner using either Ollama (local) or OpenAI\n- Website Content Extraction: Uses Playwright to render and extract content from the partner's website\n- Search Results Collection: Optionally fetches search results from Brave Search API\n- GitHub Repository Analysis: Optionally searches GitHub for repositories related to the partner")
    
    doc.add_paragraph("3. Analysis Generation:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("- Combines all collected data into a comprehensive prompt\n- Sends the prompt to an AI model (either local Ollama or OpenAI)\n- Parses the JSON response from the AI model")
    
    doc.add_paragraph("4. Results Storage:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("- Writes summary results to a CSV file\n- Saves detailed analysis for each partner as individual JSON files")
    
    # Key Components section
    doc.add_heading("Key Components", level=2)
    
    components = [
        "LLM Integration: Uses both local models via Ollama and cloud models via OpenAI API",
        "Web Scraping: Uses Playwright for JavaScript-rendered website content extraction",
        "API Integration: Connects to Brave Search and GitHub APIs for enhanced data collection",
        "Robust JSON Parsing: Implements multiple fallback mechanisms to handle various LLM response formats"
    ]
    
    for component in components:
        doc.add_paragraph(component, style="List Bullet")
    
    # Configuration Options section
    doc.add_heading("Configuration Options", level=2)
    
    doc.add_paragraph("The script offers extensive configuration through command-line arguments:")
    
    options = [
        "Input/Output Options: Specify input files, output directory, and format",
        "Model Selection: Choose between different Ollama and OpenAI models",
        "Feature Flags: Enable/disable website content fetching, search enhancement, and GitHub analysis",
        "Logging Options: Configure verbosity levels"
    ]
    
    for option in options:
        doc.add_paragraph(option, style="List Bullet")
    
    # Error Handling section
    doc.add_heading("Error Handling", level=2)
    
    doc.add_paragraph("The script implements comprehensive error handling:")
    
    handling = [
        "Catches and logs exceptions during partner processing",
        "Continues processing remaining partners if one fails",
        "Implements retry mechanisms for API rate limiting"
    ]
    
    for item in handling:
        doc.add_paragraph(item, style="List Bullet")
    
    # Conclusion
    doc.add_paragraph("This partner analysis system provides a sophisticated way to evaluate partners against AI technology categories, leveraging multiple data sources and AI models to generate comprehensive analysis results.")
    
    # Save the document
    output_path = os.path.join(output_dir, "Partner_Analysis_Summary.docx")
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    main() 