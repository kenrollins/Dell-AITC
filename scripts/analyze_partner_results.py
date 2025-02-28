#!/usr/bin/env python3
"""
Partner Analysis Results Summary Script

This script analyzes the consolidated CSV file from partner_analysis.py and generates
a comprehensive summary of the results for use in presentations.

Usage:
    python analyze_partner_results.py [options]
    
    # Generate a Word document report
    python analyze_partner_results.py --output-format word
    
    # Generate a Markdown report
    python analyze_partner_results.py --output-format markdown
    
    # Use the wrapper script (recommended)
    bash scripts/run_analysis_report.sh word
    bash scripts/run_analysis_report.sh markdown

Options:
    --input-file FILE          Path to the consolidated CSV file (default: latest partner_analysis_*.csv file)
    --output-file FILE         Path to save the summary report (default: partner_summary_{timestamp}.md or .docx)
    --output-dir DIR           Directory to save output files (default: data/output/partner_analysis)
    --output-format FORMAT     Output format: 'markdown' or 'word' (default: markdown)
    --min-confidence FLOAT     Minimum confidence threshold (default: 0.7)
    --verbose                  Enable verbose logging

Notes:
    - The Word output format requires the python-docx package to be installed
    - The Markdown output format is compatible with most Markdown viewers
    - Visualizations are saved as PNG files in the same directory as the output file
    - Output files are named partner_summary_{timestamp}.docx or partner_summary_{timestamp}.md by default
    - Input files are expected to follow the naming convention partner_analysis_{timestamp}.csv
"""

import os
import csv
import glob
import argparse
import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from collections import defaultdict, Counter
from typing import Dict, List, Any, Tuple
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def _parse_args():
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(description="Analyze partner analysis results and generate a summary.")
    
    parser.add_argument("--input-file", type=str, default=None, 
                        help="Path to the consolidated CSV file (default: latest partner_analysis_*.csv file)")
    parser.add_argument("--output-file", type=str, default=None,
                        help="Path to save the summary report (default: partner_summary_{timestamp}.md or .docx)")
    parser.add_argument("--output-dir", type=str, default="data/output/partner_analysis",
                        help="Directory to save output files")
    parser.add_argument("--output-format", type=str, choices=['markdown', 'word'], default='markdown',
                        help="Output format: 'markdown' or 'word' (default: markdown)")
    parser.add_argument("--min-confidence", type=float, default=0.7,
                        help="Minimum confidence threshold")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose logging")
    
    args = parser.parse_args()
    
    # Find the latest consolidated CSV file if not specified
    if args.input_file is None:
        csv_files = glob.glob(os.path.join(args.output_dir, "partner_analysis_*.csv"))
        if not csv_files:
            raise ValueError(f"No consolidated CSV files found in {args.output_dir}")
        args.input_file = max(csv_files, key=os.path.getctime)
        print(f"Using latest consolidated CSV file: {args.input_file}")
    
    # Generate default output file name if not specified
    if args.output_file is None:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        extension = '.docx' if args.output_format == 'word' else '.md'
        prefix = 'partner_summary_'
        args.output_file = os.path.join(args.output_dir, f"{prefix}{timestamp}{extension}")
        print(f"Summary will be saved to: {args.output_file}")
    elif args.output_format == 'word' and not args.output_file.endswith('.docx'):
        args.output_file = args.output_file + '.docx'
        print(f"Adjusted output file to: {args.output_file}")
    elif args.output_format == 'markdown' and not args.output_file.endswith('.md'):
        args.output_file = args.output_file + '.md'
        print(f"Adjusted output file to: {args.output_file}")
    
    return args

def load_data(input_file: str) -> pd.DataFrame:
    """Load data from the consolidated CSV file."""
    print(f"Loading data from {input_file}")
    
    try:
        df = pd.read_csv(input_file)
        print(f"Loaded {len(df)} rows from {input_file}")
        
        # Normalize category names to handle inconsistencies
        print("Normalizing category names...")
        # Replace "Data Integration & Deployment" with "Data Integration & Management"
        df['category'] = df['category'].replace("Data Integration & Deployment", "Data Integration & Management")
        # Remove header row if it got included in the data
        df = df[df['category'] != 'category']
        
        # Add search_results_count column if it doesn't exist (for backward compatibility)
        if 'search_results_count' not in df.columns:
            df['search_results_count'] = 10  # Default value
        
        return df
    except Exception as e:
        print(f"Error loading data: {str(e)}")
        raise

def filter_data(df: pd.DataFrame, min_confidence: float) -> pd.DataFrame:
    """Filter data based on minimum confidence threshold."""
    print(f"Filtering data with minimum confidence threshold of {min_confidence}")
    
    filtered_df = df[df['confidence'] >= min_confidence].copy()
    print(f"Filtered data contains {len(filtered_df)} rows")
    return filtered_df

def analyze_categories(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze AI technology categories."""
    print("Analyzing AI technology categories")
    
    # Count categories
    category_counts = df['category'].value_counts().to_dict()
    
    # Calculate average confidence by category
    category_avg_confidence = df.groupby('category')['confidence'].mean().to_dict()
    
    # Find top partners for each category
    category_top_partners = {}
    for category in df['category'].unique():
        category_df = df[df['category'] == category]
        top_partners = category_df.sort_values('confidence', ascending=False)['partner_name'].unique()[:3]
        category_top_partners[category] = top_partners.tolist()
    
    # Calculate category coverage (percentage of partners in each category)
    total_partners = df['partner_name'].nunique()
    category_coverage = {}
    for category in df['category'].unique():
        partners_in_category = df[df['category'] == category]['partner_name'].nunique()
        category_coverage[category] = partners_in_category / total_partners
    
    return {
        'category_counts': category_counts,
        'category_avg_confidence': category_avg_confidence,
        'category_top_partners': category_top_partners,
        'category_coverage': category_coverage
    }

def analyze_partners(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze partners."""
    print("Analyzing partners")
    
    # Count categories per partner
    partner_category_counts = df.groupby('partner_name')['category'].nunique().to_dict()
    
    # Calculate average confidence by partner
    partner_avg_confidence = df.groupby('partner_name')['confidence'].mean().to_dict()
    
    # Find top categories for each partner
    partner_top_categories = {}
    for partner in df['partner_name'].unique():
        partner_df = df[df['partner_name'] == partner]
        top_categories = partner_df.sort_values('confidence', ascending=False)['category'].tolist()
        partner_top_categories[partner] = top_categories
    
    # Find partners with highest average confidence
    top_partners_by_confidence = sorted(partner_avg_confidence.items(), key=lambda x: x[1], reverse=True)[:10]
    
    # Find partners with most categories
    top_partners_by_categories = sorted(partner_category_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    
    return {
        'partner_category_counts': partner_category_counts,
        'partner_avg_confidence': partner_avg_confidence,
        'partner_top_categories': partner_top_categories,
        'top_partners_by_confidence': top_partners_by_confidence,
        'top_partners_by_categories': top_partners_by_categories
    }

def analyze_industry_coverage(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze industry coverage."""
    print("Analyzing industry coverage")
    
    # Extract industries from target_industries column
    all_industries = []
    for industries_str in df['target_industries'].dropna():
        industries = [industry.strip() for industry in industries_str.split(',')]
        all_industries.extend(industries)
    
    # Count industries
    industry_counts = Counter(all_industries)
    
    # Map industries to categories
    industry_categories = defaultdict(list)
    for _, row in df.iterrows():
        if pd.notna(row['target_industries']):
            industries = [industry.strip() for industry in row['target_industries'].split(',')]
            for industry in industries:
                if industry not in industry_categories[row['category']]:
                    industry_categories[row['category']].append(industry)
    
    return {
        'industry_counts': industry_counts,
        'industry_categories': industry_categories
    }

def analyze_technologies(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze technologies."""
    print("Analyzing technologies")
    
    # Extract technologies from key_technologies column
    all_technologies = []
    for technologies_str in df['key_technologies'].dropna():
        technologies = [tech.strip() for tech in technologies_str.split(',')]
        all_technologies.extend(technologies)
    
    # Count technologies
    technology_counts = Counter(all_technologies)
    
    # Map technologies to categories
    technology_categories = defaultdict(list)
    for _, row in df.iterrows():
        if pd.notna(row['key_technologies']):
            technologies = [tech.strip() for tech in row['key_technologies'].split(',')]
            for tech in technologies:
                if tech not in technology_categories[row['category']]:
                    technology_categories[row['category']].append(tech)
    
    return {
        'technology_counts': technology_counts,
        'technology_categories': technology_categories
    }

def identify_gaps(df: pd.DataFrame, all_categories: List[str]) -> Dict[str, Any]:
    """Identify gaps in category coverage."""
    print("Identifying gaps in category coverage")
    
    # Find categories with no or few partners
    covered_categories = set(df['category'].unique())
    missing_categories = set(all_categories) - covered_categories
    
    # Find categories with low confidence
    low_confidence_categories = df.groupby('category')['confidence'].mean().reset_index()
    low_confidence_categories = low_confidence_categories[low_confidence_categories['confidence'] < 0.75]['category'].tolist()
    
    # Find categories with few partners
    category_partner_counts = df.groupby('category')['partner_name'].nunique().reset_index()
    category_partner_counts.columns = ['category', 'partner_count']
    few_partners_categories = category_partner_counts[category_partner_counts['partner_count'] < 3]['category'].tolist()
    
    return {
        'missing_categories': list(missing_categories),
        'low_confidence_categories': low_confidence_categories,
        'few_partners_categories': few_partners_categories
    }

def generate_visualizations(df: pd.DataFrame, output_dir: str) -> Dict[str, str]:
    """Generate visualizations for the summary."""
    print("Generating visualizations")
    
    os.makedirs(output_dir, exist_ok=True)
    visualization_paths = {}
    
    # Set the style
    sns.set(style="whitegrid")
    
    # 1. Category distribution
    plt.figure(figsize=(12, 8))
    category_counts = df['category'].value_counts()
    sns.barplot(x=category_counts.values, y=category_counts.index)
    plt.title('Number of Partners per AI Technology Category')
    plt.xlabel('Number of Partners')
    plt.tight_layout()
    category_dist_path = os.path.join(output_dir, 'category_distribution.png')
    plt.savefig(category_dist_path)
    plt.close()
    visualization_paths['category_distribution'] = category_dist_path
    
    # 2. Average confidence by category
    plt.figure(figsize=(12, 8))
    avg_confidence = df.groupby('category')['confidence'].mean().sort_values(ascending=False)
    sns.barplot(x=avg_confidence.values, y=avg_confidence.index)
    plt.title('Average Confidence Score by AI Technology Category')
    plt.xlabel('Average Confidence')
    plt.xlim(0, 1)
    plt.tight_layout()
    avg_conf_path = os.path.join(output_dir, 'avg_confidence_by_category.png')
    plt.savefig(avg_conf_path)
    plt.close()
    visualization_paths['avg_confidence_by_category'] = avg_conf_path
    
    # 3. Top partners by number of categories
    plt.figure(figsize=(12, 8))
    partner_categories = df.groupby('partner_name')['category'].nunique().sort_values(ascending=False).head(10)
    sns.barplot(x=partner_categories.values, y=partner_categories.index)
    plt.title('Top 10 Partners by Number of AI Technology Categories')
    plt.xlabel('Number of Categories')
    plt.tight_layout()
    top_partners_path = os.path.join(output_dir, 'top_partners_by_categories.png')
    plt.savefig(top_partners_path)
    plt.close()
    visualization_paths['top_partners_by_categories'] = top_partners_path
    
    # 4. Industry distribution
    plt.figure(figsize=(12, 8))
    all_industries = []
    for industries_str in df['target_industries'].dropna():
        industries = [industry.strip() for industry in industries_str.split(',')]
        all_industries.extend(industries)
    
    industry_counts = Counter(all_industries).most_common(10)
    sns.barplot(x=[count for _, count in industry_counts], y=[industry for industry, _ in industry_counts])
    plt.title('Top 10 Industries Served by Partners')
    plt.xlabel('Number of Mentions')
    plt.tight_layout()
    industry_dist_path = os.path.join(output_dir, 'industry_distribution.png')
    plt.savefig(industry_dist_path)
    plt.close()
    visualization_paths['industry_distribution'] = industry_dist_path
    
    # 5. Technology distribution
    plt.figure(figsize=(12, 8))
    all_technologies = []
    for technologies_str in df['key_technologies'].dropna():
        technologies = [tech.strip() for tech in technologies_str.split(',')]
        all_technologies.extend(technologies)
    
    technology_counts = Counter(all_technologies).most_common(10)
    sns.barplot(x=[count for _, count in technology_counts], y=[tech for tech, _ in technology_counts])
    plt.title('Top 10 Technologies Used by Partners')
    plt.xlabel('Number of Mentions')
    plt.tight_layout()
    tech_dist_path = os.path.join(output_dir, 'technology_distribution.png')
    plt.savefig(tech_dist_path)
    plt.close()
    visualization_paths['technology_distribution'] = tech_dist_path
    
    # 6. Heatmap of categories vs industries
    # Create a matrix of categories vs top industries
    top_industries = [industry for industry, _ in Counter(all_industries).most_common(10)]
    category_industry_matrix = pd.DataFrame(0, index=df['category'].unique(), columns=top_industries)
    
    for _, row in df.iterrows():
        if pd.notna(row['target_industries']):
            industries = [industry.strip() for industry in row['target_industries'].split(',')]
            for industry in industries:
                if industry in top_industries:
                    category_industry_matrix.loc[row['category'], industry] += 1
    
    plt.figure(figsize=(14, 10))
    sns.heatmap(category_industry_matrix, annot=True, cmap='YlGnBu', fmt='d')
    plt.title('AI Technology Categories vs Top Industries')
    plt.tight_layout()
    heatmap_path = os.path.join(output_dir, 'category_industry_heatmap.png')
    plt.savefig(heatmap_path)
    plt.close()
    visualization_paths['category_industry_heatmap'] = heatmap_path
    
    return visualization_paths

def generate_summary_report(
    df: pd.DataFrame,
    category_analysis: Dict[str, Any],
    partner_analysis: Dict[str, Any],
    industry_analysis: Dict[str, Any],
    technology_analysis: Dict[str, Any],
    gap_analysis: Dict[str, Any],
    visualization_paths: Dict[str, str],
    output_file: str
) -> None:
    """Generate a comprehensive summary report."""
    print(f"Generating summary report to {output_file}")
    
    # Create the directory if it doesn't exist
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    with open(output_file, 'w') as f:
        # Title and introduction
        f.write("# Partner Analysis Summary Report\n\n")
        f.write(f"*Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n")
        
        f.write("## Executive Summary\n\n")
        
        total_partners = df['partner_name'].nunique()
        total_categories = df['category'].nunique()
        avg_confidence = df['confidence'].mean()
        
        f.write(f"This report analyzes {total_partners} partners across {total_categories} AI technology categories. ")
        f.write(f"The average confidence score across all matches is {avg_confidence:.2f}.\n\n")
        
        # Highlight key findings
        f.write("### Key Findings\n\n")
        
        # Top categories by partner count
        top_categories = sorted(category_analysis['category_counts'].items(), key=lambda x: x[1], reverse=True)[:3]
        f.write("**Top Categories by Partner Count:**\n")
        for category, count in top_categories:
            coverage = category_analysis['category_coverage'][category] * 100
            f.write(f"- {category}: {count} partners ({coverage:.1f}% of all partners)\n")
        f.write("\n")
        
        # Categories with highest average confidence
        top_confidence = sorted(category_analysis['category_avg_confidence'].items(), key=lambda x: x[1], reverse=True)[:3]
        f.write("**Categories with Highest Average Confidence:**\n")
        for category, confidence in top_confidence:
            f.write(f"- {category}: {confidence:.2f}\n")
        f.write("\n")
        
        # Top partners by category count
        f.write("**Top Partners by Category Coverage:**\n")
        for partner, count in partner_analysis['top_partners_by_categories'][:3]:
            f.write(f"- {partner}: {count} categories\n")
        f.write("\n")
        
        # Gaps in coverage
        f.write("**Gaps in Coverage:**\n")
        if gap_analysis['missing_categories']:
            f.write("- Missing categories: " + ", ".join(gap_analysis['missing_categories']) + "\n")
        if gap_analysis['few_partners_categories']:
            f.write("- Categories with few partners: " + ", ".join(gap_analysis['few_partners_categories']) + "\n")
        if gap_analysis['low_confidence_categories']:
            f.write("- Categories with low confidence: " + ", ".join(gap_analysis['low_confidence_categories']) + "\n")
        f.write("\n")
        
        # Category Analysis
        f.write("## AI Technology Category Analysis\n\n")
        
        # Category distribution
        f.write("### Category Distribution\n\n")
        f.write(f"![Category Distribution]({os.path.basename(visualization_paths['category_distribution'])})\n\n")
        
        # Average confidence by category
        f.write("### Average Confidence by Category\n\n")
        f.write(f"![Average Confidence by Category]({os.path.basename(visualization_paths['avg_confidence_by_category'])})\n\n")
        
        # Detailed category analysis
        f.write("### Detailed Category Analysis\n\n")
        
        f.write("| Category | Partner Count | Avg Confidence | Top Partners |\n")
        f.write("|----------|---------------|----------------|-------------|\n")
        
        for category in sorted(category_analysis['category_counts'].keys()):
            count = category_analysis['category_counts'][category]
            confidence = category_analysis['category_avg_confidence'][category]
            top_partners = ", ".join(category_analysis['category_top_partners'][category])
            
            f.write(f"| {category} | {count} | {confidence:.2f} | {top_partners} |\n")
        
        f.write("\n")
        
        # Partner Analysis
        f.write("## Partner Analysis\n\n")
        
        # Top partners by categories
        f.write("### Top Partners by Number of Categories\n\n")
        f.write(f"![Top Partners by Categories]({os.path.basename(visualization_paths['top_partners_by_categories'])})\n\n")
        
        # Top partners table
        f.write("### Top 10 Partners\n\n")
        
        f.write("| Partner | Categories | Avg Confidence | Top Categories |\n")
        f.write("|---------|------------|----------------|----------------|\n")
        
        for partner, count in partner_analysis['top_partners_by_categories'][:10]:
            confidence = partner_analysis['partner_avg_confidence'][partner]
            top_categories = ", ".join(partner_analysis['partner_top_categories'][partner][:3])
            
            f.write(f"| {partner} | {count} | {confidence:.2f} | {top_categories} |\n")
        
        f.write("\n")
        
        # Industry Analysis
        f.write("## Industry Analysis\n\n")
        
        # Industry distribution
        f.write("### Top Industries Served\n\n")
        f.write(f"![Industry Distribution]({os.path.basename(visualization_paths['industry_distribution'])})\n\n")
        
        # Category-Industry heatmap
        f.write("### AI Technology Categories vs Industries\n\n")
        f.write(f"![Category-Industry Heatmap]({os.path.basename(visualization_paths['category_industry_heatmap'])})\n\n")
        
        # Top industries by mention
        f.write("### Top Industries by Mention\n\n")
        
        f.write("| Industry | Mentions | Top Categories |\n")
        f.write("|----------|----------|----------------|\n")
        
        for industry, count in industry_analysis['industry_counts'].most_common(10):
            # Find categories that mention this industry
            categories = []
            for category, industries in industry_analysis['industry_categories'].items():
                if industry in industries:
                    categories.append(category)
            
            top_categories = ", ".join(categories[:3])
            
            f.write(f"| {industry} | {count} | {top_categories} |\n")
        
        f.write("\n")
        
        # Technology Analysis
        f.write("## Technology Analysis\n\n")
        
        # Technology distribution
        f.write("### Top Technologies Used\n\n")
        f.write(f"![Technology Distribution]({os.path.basename(visualization_paths['technology_distribution'])})\n\n")
        
        # Top technologies by mention
        f.write("### Top Technologies by Mention\n\n")
        
        f.write("| Technology | Mentions | Top Categories |\n")
        f.write("|------------|----------|----------------|\n")
        
        for technology, count in technology_analysis['technology_counts'].most_common(10):
            # Find categories that mention this technology
            categories = []
            for category, technologies in technology_analysis['technology_categories'].items():
                if technology in technologies:
                    categories.append(category)
            
            top_categories = ", ".join(categories[:3])
            
            f.write(f"| {technology} | {count} | {top_categories} |\n")
        
        f.write("\n")
        
        # Gap Analysis
        f.write("## Gap Analysis\n\n")
        
        # Missing categories
        if gap_analysis['missing_categories']:
            f.write("### Missing Categories\n\n")
            f.write("The following AI technology categories have no partners:\n\n")
            for category in gap_analysis['missing_categories']:
                f.write(f"- {category}\n")
            f.write("\n")
        
        # Categories with few partners
        if gap_analysis['few_partners_categories']:
            f.write("### Categories with Few Partners\n\n")
            f.write("The following categories have fewer than 3 partners:\n\n")
            for category in gap_analysis['few_partners_categories']:
                count = df[df['category'] == category]['partner_name'].nunique()
                f.write(f"- {category}: {count} partners\n")
            f.write("\n")
        
        # Categories with low confidence
        if gap_analysis['low_confidence_categories']:
            f.write("### Categories with Low Confidence\n\n")
            f.write("The following categories have an average confidence score below 0.75:\n\n")
            for category in gap_analysis['low_confidence_categories']:
                confidence = df[df['category'] == category]['confidence'].mean()
                f.write(f"- {category}: {confidence:.2f} average confidence\n")
            f.write("\n")
        
        # Recommendations
        f.write("## Recommendations\n\n")
        
        f.write("Based on the analysis, here are some recommendations:\n\n")
        
        # Recommendations for missing categories
        if gap_analysis['missing_categories'] or gap_analysis['few_partners_categories']:
            f.write("### Address Category Gaps\n\n")
            f.write("Consider identifying and onboarding partners in these underrepresented categories:\n\n")
            for category in gap_analysis['missing_categories'] + gap_analysis['few_partners_categories']:
                f.write(f"- {category}\n")
            f.write("\n")
        
        # Recommendations for low confidence categories
        if gap_analysis['low_confidence_categories']:
            f.write("### Strengthen Low Confidence Categories\n\n")
            f.write("For categories with low confidence scores, consider:\n\n")
            f.write("- Conducting deeper analysis of partner capabilities\n")
            f.write("- Providing additional training or resources to partners\n")
            f.write("- Refining category definitions for better matching\n\n")
        
        # Recommendations for industry focus
        f.write("### Industry Focus\n\n")
        top_industries = [industry for industry, _ in industry_analysis['industry_counts'].most_common(5)]
        f.write(f"The analysis shows strong representation in {', '.join(top_industries)}. ")
        f.write("Consider whether this aligns with strategic priorities or if adjustments are needed.\n\n")
        
        # Recommendations for technology focus
        f.write("### Technology Focus\n\n")
        top_technologies = [tech for tech, _ in technology_analysis['technology_counts'].most_common(5)]
        f.write(f"The analysis shows strong representation in {', '.join(top_technologies)} technologies. ")
        f.write("Evaluate if these align with market trends and future technology directions.\n\n")
        
        # Conclusion
        f.write("## Conclusion\n\n")
        
        f.write("This analysis provides a comprehensive overview of the partner ecosystem across AI technology categories. ")
        f.write("It highlights areas of strength, identifies gaps, and offers recommendations for strategic focus.\n\n")
        
        f.write("The data suggests that the partner ecosystem is particularly strong in ")
        if top_categories and isinstance(top_categories[0], tuple):
            f.write(", ".join([category for category, _ in top_categories]))
        else:
            f.write(", ".join([str(category) for category in top_categories]))
        f.write(". ")
        
        if gap_analysis['missing_categories'] or gap_analysis['few_partners_categories']:
            f.write("However, there are opportunities to expand coverage in ")
            gaps = gap_analysis['missing_categories'] + gap_analysis['few_partners_categories']
            f.write(", ".join(gaps[:3]))
            if len(gaps) > 3:
                f.write(f", and {len(gaps) - 3} other categories")
            f.write(".\n\n")
        else:
            f.write("The partner ecosystem provides good coverage across all AI technology categories.\n\n")
        
        f.write("By addressing the recommendations in this report, the partner ecosystem can be strengthened ")
        f.write("to better align with strategic priorities and market opportunities.\n")
    
    print(f"Summary report generated: {output_file}")

def generate_word_report(
    df: pd.DataFrame,
    category_analysis: Dict[str, Any],
    partner_analysis: Dict[str, Any],
    industry_analysis: Dict[str, Any],
    technology_analysis: Dict[str, Any],
    gap_analysis: Dict[str, Any],
    visualization_paths: Dict[str, str],
    output_file: str
) -> None:
    """Generate a comprehensive summary report in Word format."""
    print(f"Generating Word summary report to {output_file}")
    
    # Create the directory if it doesn't exist
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Add styles for headings and text
    styles = doc.styles
    
    # Title and introduction
    doc.add_heading("Partner Analysis Summary Report", level=0)
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    
    total_partners = df['partner_name'].nunique()
    total_categories = df['category'].nunique()
    avg_confidence = df['confidence'].mean()
    
    summary_text = f"This report analyzes {total_partners} partners across {total_categories} AI technology categories. "
    summary_text += f"The average confidence score across all matches is {avg_confidence:.2f}."
    doc.add_paragraph(summary_text)
    
    # Key Findings
    doc.add_heading("Key Findings", level=2)
    
    # Top categories by partner count
    p = doc.add_paragraph()
    p.add_run("Top Categories by Partner Count:").bold = True
    top_categories = sorted(category_analysis['category_counts'].items(), key=lambda x: x[1], reverse=True)[:3]
    for category, count in top_categories:
        coverage = category_analysis['category_coverage'][category] * 100
        doc.add_paragraph(f"{category}: {count} partners ({coverage:.1f}% of all partners)", style='List Bullet')
    
    # Categories with highest average confidence
    p = doc.add_paragraph()
    p.add_run("Categories with Highest Average Confidence:").bold = True
    top_confidence = sorted(category_analysis['category_avg_confidence'].items(), key=lambda x: x[1], reverse=True)[:3]
    for category, confidence in top_confidence:
        doc.add_paragraph(f"{category}: {confidence:.2f}", style='List Bullet')
    
    # Top partners by category count
    p = doc.add_paragraph()
    p.add_run("Top Partners by Category Coverage:").bold = True
    for partner, count in partner_analysis['top_partners_by_categories'][:3]:
        doc.add_paragraph(f"{partner}: {count} categories", style='List Bullet')
    
    # Gaps in coverage
    p = doc.add_paragraph()
    p.add_run("Gaps in Coverage:").bold = True
    if gap_analysis['missing_categories']:
        doc.add_paragraph("Missing categories: " + ", ".join(gap_analysis['missing_categories']), style='List Bullet')
    if gap_analysis['few_partners_categories']:
        doc.add_paragraph("Categories with few partners: " + ", ".join(gap_analysis['few_partners_categories']), style='List Bullet')
    if gap_analysis['low_confidence_categories']:
        doc.add_paragraph("Categories with low confidence: " + ", ".join(gap_analysis['low_confidence_categories']), style='List Bullet')
    
    # Category Analysis
    doc.add_heading("AI Technology Category Analysis", level=1)
    
    # Category distribution
    doc.add_heading("Category Distribution", level=2)
    doc.add_picture(visualization_paths['category_distribution'], width=Inches(6))
    
    # Average confidence by category
    doc.add_heading("Average Confidence by Category", level=2)
    doc.add_picture(visualization_paths['avg_confidence_by_category'], width=Inches(6))
    
    # Detailed category analysis
    doc.add_heading("Detailed Category Analysis", level=2)
    
    # Create a table for category analysis
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Partner Count"
    header_cells[2].text = "Avg Confidence"
    header_cells[3].text = "Top Partners"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for category in sorted(category_analysis['category_counts'].keys()):
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(category_analysis['category_counts'][category])
        row_cells[2].text = f"{category_analysis['category_avg_confidence'][category]:.2f}"
        row_cells[3].text = ", ".join(category_analysis['category_top_partners'][category])
    
    doc.add_paragraph()  # Add space after table
    
    # Partner Analysis
    doc.add_heading("Partner Analysis", level=1)
    
    # Top partners by categories
    doc.add_heading("Top Partners by Number of Categories", level=2)
    doc.add_picture(visualization_paths['top_partners_by_categories'], width=Inches(6))
    
    # Top partners table
    doc.add_heading("Top 10 Partners", level=2)
    
    # Create a table for partner analysis
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Partner"
    header_cells[1].text = "Categories"
    header_cells[2].text = "Avg Confidence"
    header_cells[3].text = "Top Categories"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for partner, count in partner_analysis['top_partners_by_categories'][:10]:
        row_cells = table.add_row().cells
        row_cells[0].text = partner
        row_cells[1].text = str(count)
        row_cells[2].text = f"{partner_analysis['partner_avg_confidence'][partner]:.2f}"
        row_cells[3].text = ", ".join(partner_analysis['partner_top_categories'][partner][:3])
    
    doc.add_paragraph()  # Add space after table
    
    # Industry Analysis
    doc.add_heading("Industry Analysis", level=1)
    
    # Industry distribution
    doc.add_heading("Top Industries Served", level=2)
    doc.add_picture(visualization_paths['industry_distribution'], width=Inches(6))
    
    # Category-Industry heatmap
    doc.add_heading("AI Technology Categories vs Industries", level=2)
    doc.add_picture(visualization_paths['category_industry_heatmap'], width=Inches(6))
    
    # Top industries by mention
    doc.add_heading("Top Industries by Mention", level=2)
    
    # Create a table for industry analysis
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Industry"
    header_cells[1].text = "Mentions"
    header_cells[2].text = "Top Categories"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for industry, count in industry_analysis['industry_counts'].most_common(10):
        # Find categories that mention this industry
        categories = []
        for category, industries in industry_analysis['industry_categories'].items():
            if industry in industries:
                categories.append(category)
        
        top_categories = ", ".join(categories[:3])
        
        row_cells = table.add_row().cells
        row_cells[0].text = industry
        row_cells[1].text = str(count)
        row_cells[2].text = top_categories
    
    doc.add_paragraph()  # Add space after table
    
    # Technology Analysis
    doc.add_heading("Technology Analysis", level=1)
    
    # Technology distribution
    doc.add_heading("Top Technologies Used", level=2)
    doc.add_picture(visualization_paths['technology_distribution'], width=Inches(6))
    
    # Top technologies by mention
    doc.add_heading("Top Technologies by Mention", level=2)
    
    # Create a table for technology analysis
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Technology"
    header_cells[1].text = "Mentions"
    header_cells[2].text = "Top Categories"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for technology, count in technology_analysis['technology_counts'].most_common(10):
        # Find categories that mention this technology
        categories = []
        for category, technologies in technology_analysis['technology_categories'].items():
            if technology in technologies:
                categories.append(category)
        
        top_categories = ", ".join(categories[:3])
        
        row_cells = table.add_row().cells
        row_cells[0].text = technology
        row_cells[1].text = str(count)
        row_cells[2].text = top_categories
    
    doc.add_paragraph()  # Add space after table
    
    # Gap Analysis
    doc.add_heading("Gap Analysis", level=1)
    
    # Missing categories
    if gap_analysis['missing_categories']:
        doc.add_heading("Missing Categories", level=2)
        doc.add_paragraph("The following AI technology categories have no partners:")
        for category in gap_analysis['missing_categories']:
            doc.add_paragraph(category, style='List Bullet')
    
    # Categories with few partners
    if gap_analysis['few_partners_categories']:
        doc.add_heading("Categories with Few Partners", level=2)
        doc.add_paragraph("The following categories have fewer than 3 partners:")
        for category in gap_analysis['few_partners_categories']:
            count = df[df['category'] == category]['partner_name'].nunique()
            doc.add_paragraph(f"{category}: {count} partners", style='List Bullet')
    
    # Categories with low confidence
    if gap_analysis['low_confidence_categories']:
        doc.add_heading("Categories with Low Confidence", level=2)
        doc.add_paragraph("The following categories have an average confidence score below 0.75:")
        for category in gap_analysis['low_confidence_categories']:
            confidence = df[df['category'] == category]['confidence'].mean()
            doc.add_paragraph(f"{category}: {confidence:.2f} average confidence", style='List Bullet')
    
    # Recommendations
    doc.add_heading("Recommendations", level=1)
    
    doc.add_paragraph("Based on the analysis, here are some recommendations:")
    
    # Recommendations for missing categories
    if gap_analysis['missing_categories'] or gap_analysis['few_partners_categories']:
        doc.add_heading("Address Category Gaps", level=2)
        doc.add_paragraph("Consider identifying and onboarding partners in these underrepresented categories:")
        for category in gap_analysis['missing_categories'] + gap_analysis['few_partners_categories']:
            doc.add_paragraph(category, style='List Bullet')
    
    # Recommendations for low confidence categories
    if gap_analysis['low_confidence_categories']:
        doc.add_heading("Strengthen Low Confidence Categories", level=2)
        doc.add_paragraph("For categories with low confidence scores, consider:")
        doc.add_paragraph("Conducting deeper analysis of partner capabilities", style='List Bullet')
        doc.add_paragraph("Providing additional training or resources to partners", style='List Bullet')
        doc.add_paragraph("Refining category definitions for better matching", style='List Bullet')
    
    # Recommendations for industry focus
    doc.add_heading("Industry Focus", level=2)
    top_industries = [industry for industry, _ in industry_analysis['industry_counts'].most_common(5)]
    industry_text = f"The analysis shows strong representation in {', '.join(top_industries)}. "
    industry_text += "Consider whether this aligns with strategic priorities or if adjustments are needed."
    doc.add_paragraph(industry_text)
    
    # Recommendations for technology focus
    doc.add_heading("Technology Focus", level=2)
    top_technologies = [tech for tech, _ in technology_analysis['technology_counts'].most_common(5)]
    tech_text = f"The analysis shows strong representation in {', '.join(top_technologies)} technologies. "
    tech_text += "Evaluate if these align with market trends and future technology directions."
    doc.add_paragraph(tech_text)
    
    # Conclusion
    doc.add_heading("Conclusion", level=1)
    
    conclusion_text = "This analysis provides a comprehensive overview of the partner ecosystem across AI technology categories. "
    conclusion_text += "It highlights areas of strength, identifies gaps, and offers recommendations for strategic focus.\n\n"
    conclusion_text += "The data suggests that the partner ecosystem is particularly strong in "
    
    if top_categories and isinstance(top_categories[0], tuple):
        conclusion_text += ", ".join([category for category, _ in top_categories])
    else:
        conclusion_text += ", ".join([str(category) for category in top_categories])
    conclusion_text += ". "
    
    if gap_analysis['missing_categories'] or gap_analysis['few_partners_categories']:
        conclusion_text += "However, there are opportunities to expand coverage in "
        gaps = gap_analysis['missing_categories'] + gap_analysis['few_partners_categories']
        conclusion_text += ", ".join(gaps[:3])
        if len(gaps) > 3:
            conclusion_text += f", and {len(gaps) - 3} other categories"
        conclusion_text += "."
    else:
        conclusion_text += "The partner ecosystem provides good coverage across all AI technology categories."
    
    doc.add_paragraph(conclusion_text)
    
    final_text = "By addressing the recommendations in this report, the partner ecosystem can be strengthened "
    final_text += "to better align with strategic priorities and market opportunities."
    doc.add_paragraph(final_text)
    
    # Save the document
    doc.save(output_file)
    
    print(f"Word summary report generated: {output_file}")

def main():
    """Main function."""
    print("Starting partner analysis results summary")
    
    # Parse command-line arguments
    args = _parse_args()
    
    # Load data
    df = load_data(args.input_file)
    
    # Filter data
    filtered_df = filter_data(df, args.min_confidence)
    
    # Define all possible categories (this should ideally come from a reference file)
    # For now, we'll use the categories found in the data
    all_categories = filtered_df['category'].unique().tolist()
    
    # Analyze categories
    category_analysis = analyze_categories(filtered_df)
    
    # Analyze partners
    partner_analysis = analyze_partners(filtered_df)
    
    # Analyze industry coverage
    industry_analysis = analyze_industry_coverage(filtered_df)
    
    # Analyze technologies
    technology_analysis = analyze_technologies(filtered_df)
    
    # Identify gaps
    gap_analysis = identify_gaps(filtered_df, all_categories)
    
    # Generate visualizations
    visualization_paths = generate_visualizations(filtered_df, os.path.dirname(args.output_file))
    
    # Generate summary report based on output format
    if args.output_format == 'word':
        generate_word_report(
            filtered_df,
            category_analysis,
            partner_analysis,
            industry_analysis,
            technology_analysis,
            gap_analysis,
            visualization_paths,
            args.output_file
        )
    else:  # markdown
        generate_summary_report(
            filtered_df,
            category_analysis,
            partner_analysis,
            industry_analysis,
            technology_analysis,
            gap_analysis,
            visualization_paths,
            args.output_file
        )
    
    print(f"Summary report generated: {args.output_file}")

if __name__ == "__main__":
    main() 